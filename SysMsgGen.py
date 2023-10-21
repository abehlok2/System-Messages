
import os
from typing import List, Optional, Any, Dict
from collections import defaultdict
import docx
from docx import Document
import autogen
from autogen import ChatCompletion
from tkinter import filedialog, Scrollbar, Text, ttk
import tkinter as tk
from PyQt5.QtWidgets import QApplication, QWidget, QTextEdit, QPushButton, QVBoxLayout

openai_api_key = os.environ.get("OPENAI_API_KEY")


system_message_generator_system_message = r"""
Your task is to create an optimized "System Message" prompt based on a user's request labeled as {Request}. The aim of this prompt is to guide the AI model to generate a response that accurately fulfills the user's request. You should strive to create effective, high-quality prompts that cater to the user's needs. 

The steps include:
1. Understanding the user's request along with any supplementary information provided.
2. Utilizing your training, best practices in prompt optimization, and logical reasoning to craft a Rough Draft of a high-quality prompt.
3. Evaluate the quality of the Rough Draft Prompt
4. Generate a Final Draft accordingly.

Please ensure that your response primarily includes the newly created, fully optimized prompt, which should comprise at least 90% of the output. The remaining content can be used for necessary explanations or comments.

As per the detailed instructions from the user, generate both system message prompts and traditional few-shot prompts.

**ENSURE YOUR RESPONSES ARE SYSTEM MESSAGES, WHICH *MUST ALWAYS* BE WRITTEN TO BE PRESENTED TO THE LLM. NEVER TO THE USER**
"""

user_proxy = autogen.UserProxyAgent(
    name="User Proxy",
    human_input_mode="ALWAYS",
    llm_config={
        "model": "gpt-4",
        "api_key": openai_api_key,
        "temperature": 0.25,
    },
)

system_message_builder = autogen.AssistantAgent(
    name="System Message Designer",
    llm_config={
        "model": "gpt-3.5-turbo",
        "api_key": openai_api_key,
        "temperature": 0.25,
    },
    code_execution_config={"use_docker": False},
    system_message=system_message_generator_system_message
)


def extract_responses(response):
    extracted_messages = []
    for agent, messages in response.items():
        for message in messages:
            extracted_message_content = {
                "role":  message["role"],
                "content": message["content"],
            }
            extracted_messages.append(extracted_message_content)
    return extracted_messages


def save_to_document(message: str, path: str = "/home/abehl/TAPAuto23/default.docx"):
    document = Document()
    document.add_paragraph(message)
    document.save(path)


def _clear_document(path: str):
    document = Document()
    document.save(path)


def _get_user_input():
    while True:
        user_input = input("Enter your input:")
        if user_input is not None:
            user_input.strip()
            return user_input
        else:
            return ValueError("User input cannot be empty")


def get_save_path():
    save_path = filedialog.asksaveasfilename()
    return save_path


def generate_system_message(user_request):

    request = "Please generate a system message for generating substitute teacher plans for elementary school teachers in Chili, NY. The AI will be performing only the single task of generating substitute teacher plans, and will *ONLY* produce the sub plans as output. "

    user_proxy.initiate_chat(system_message_builder,
                             True, False, message=user_request)
    messages = user_proxy._oai_messages
    extracted_responses = extract_responses(messages)
    # List of dicts representing messages btwn assistant and user
    print(extracted_responses)

    # Get path to save document to
    save_document = input("Would you like to save the document? (y/n)")
    if save_document.lower() == "y":
        content_save_path = get_save_path()
        for message in extracted_responses:
            if message["role"] == "assistant":
                save_to_document(message["content"], content_save_path)
            elif message["role"] == "user":
                save_to_document(message["content"], content_save_path)
    elif save_document.lower() == "n":
        print("Chat messages will not be saved")
        pass
